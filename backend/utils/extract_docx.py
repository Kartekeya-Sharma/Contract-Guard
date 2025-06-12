from docx import Document

def extract_text_from_docx(filepath):
    """
    Extract text from a DOCX file using python-docx
    """
    try:
        doc = Document(filepath)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Skip empty cells
                        full_text.append(cell.text)
        
        # Join all text with newlines
        text = '\n'.join(full_text)
        
        # Clean up the text
        text = text.replace('\n\n', '\n')  # Remove double newlines
        text = ' '.join(text.split())  # Normalize whitespace
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}") 